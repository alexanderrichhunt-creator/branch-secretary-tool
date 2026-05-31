from __future__ import annotations

import io
from datetime import date, datetime
from pathlib import Path

from .hymns import (
    HYMN_BOOK_CHILDREN,
    HYMN_BOOK_HYMNS,
    hymn_book_label,
    hymn_display,
    hymn_lyrics,
    hymn_title,
    normalize_hymn_book,
    parse_hymn_number,
)

_COVER_IMAGE_PATH = Path(__file__).resolve().parent / "static" / "images" / "baptism_cover.jpg"

DEFAULT_BAPTISM = {
    "service_date": "",
    "service_time": "",
    "location": "",
    "presiding": "",
    "conducting": "",
    "welcome_text": "We welcome you to this baptismal service.",
    "opening_hymn_num": "2",
    "opening_hymn_title": "",
    "opening_hymn_book": HYMN_BOOK_CHILDREN,
    "invocation": "(by invitation)",
    "speaker_1": "",
    "speaker_1_topic": "",
    "speaker_2": "",
    "speaker_2_topic": "",
    "musical_number": "",
    "candidate_name": "",
    "baptism_by": "",
    "confirmation_by": "",
    "confirmation_text": (
        "Following the baptism, [candidate] will be confirmed a member of The Church of Jesus Christ "
        "of Latter-day Saints and receive the gift of the Holy Ghost."
    ),
    "closing_hymn_num": "120",
    "closing_hymn_title": "",
    "closing_hymn_book": HYMN_BOOK_CHILDREN,
    "benediction": "(by invitation)",
    "reception_notes": "",
}

SAVABLE_BAPTISM_KEYS = (
    "presiding",
    "conducting",
    "welcome_text",
    "opening_hymn_num",
    "opening_hymn_title",
    "opening_hymn_book",
    "invocation",
    "speaker_1",
    "speaker_1_topic",
    "speaker_2",
    "speaker_2_topic",
    "musical_number",
    "confirmation_text",
    "closing_hymn_num",
    "closing_hymn_title",
    "closing_hymn_book",
    "benediction",
    "reception_notes",
    "location",
)


def get_branch_baptism_defaults() -> dict:
    from .models import BaptismDefaults

    merged = dict(DEFAULT_BAPTISM)
    row = BaptismDefaults.query.get(1)
    if not row:
        return merged
    for key in SAVABLE_BAPTISM_KEYS:
        merged[key] = getattr(row, key) or ""
    return merged


def has_saved_baptism_defaults() -> bool:
    from .models import BaptismDefaults

    return BaptismDefaults.query.get(1) is not None


def save_branch_baptism_defaults(form) -> None:
    from . import db
    from .models import BaptismDefaults

    row = BaptismDefaults.query.get(1)
    if not row:
        row = BaptismDefaults(id=1)
        db.session.add(row)
    for key in SAVABLE_BAPTISM_KEYS:
        value = (form.get(key) or "").strip()
        if key == "confirmation_text":
            value = confirmation_text_to_template(value, form.get("candidate_name"))
        setattr(row, key, value)
    row.updated_at = datetime.utcnow()
    db.session.commit()


def effective_hymn_title(num_raw: str | None, title_raw: str | None, book_raw: str | None) -> str:
    title = (title_raw or "").strip()
    if title:
        return title
    number = parse_hymn_number(num_raw)
    if not number:
        return ""
    return hymn_title(number, normalize_hymn_book(book_raw))


def resolved_hymn_title(defaults: dict, num_key: str, title_key: str, book_key: str) -> str:
    return effective_hymn_title(
        defaults.get(num_key),
        defaults.get(title_key),
        defaults.get(book_key),
    )


def confirmation_text_to_template(text: str | None, candidate: str | None = None) -> str:
    """Normalize saved confirmation remarks to a template with [candidate]."""
    text = (text or "").strip()
    if not text:
        return DEFAULT_BAPTISM["confirmation_text"]
    if "[candidate]" in text:
        return text
    candidate = (candidate or "").strip()
    if candidate and candidate in text:
        return text.replace(candidate, "[candidate]")
    if "receive the gift of the Holy Ghost" in text and "Following the baptism" in text:
        return DEFAULT_BAPTISM["confirmation_text"]
    return text


def resolve_confirmation_text(candidate: str | None, template: str | None) -> str:
    text = confirmation_text_to_template(template, candidate)
    candidate = (candidate or "").strip()
    if not candidate:
        return text
    if "[candidate]" in text:
        return text.replace("[candidate]", candidate)
    if candidate.lower() in text.lower():
        return text
    return text


def _hymn_export_fields(form, prefix: str) -> dict:
    num_raw = (form.get(f"{prefix}_hymn_num") or "").strip()
    book = normalize_hymn_book(form.get(f"{prefix}_hymn_book"))
    title = effective_hymn_title(num_raw, form.get(f"{prefix}_hymn_title"), book)
    number = parse_hymn_number(num_raw)
    label = hymn_book_label(book) if book == HYMN_BOOK_CHILDREN else None
    lyrics = hymn_lyrics(number, book)
    return {
        "num_raw": num_raw,
        "number": number,
        "book": book,
        "title": title,
        "line": hymn_display(num_raw, title, book_label=label),
        "lyrics": hymn_lyrics(number, book),
    }


def _format_service_date(d: date | str | None) -> str:
    if isinstance(d, str):
        try:
            d = datetime.strptime(d, "%Y-%m-%d").date()
        except ValueError:
            return d
    if not d:
        return ""
    return d.strftime("%A, %B %d, %Y").replace(" 0", " ")


def _format_service_time(raw: str | None) -> str:
    raw = (raw or "").strip()
    if not raw:
        return ""
    try:
        t = datetime.strptime(raw, "%H:%M").time()
        return t.strftime("%I:%M %p").lstrip("0")
    except ValueError:
        return raw


def baptism_from_form(form) -> dict:
    service_date_raw = (form.get("service_date") or "").strip()
    service_date = None
    if service_date_raw:
        try:
            service_date = datetime.strptime(service_date_raw, "%Y-%m-%d").date()
        except ValueError:
            service_date = None

    opening = _hymn_export_fields(form, "opening")
    closing = _hymn_export_fields(form, "closing")

    candidate = (form.get("candidate_name") or "").strip()
    confirmation_template = confirmation_text_to_template(form.get("confirmation_text"), candidate)
    confirmation_text = resolve_confirmation_text(candidate, confirmation_template)

    return {
        "service_date": service_date,
        "service_date_display": _format_service_date(service_date or service_date_raw),
        "service_time_display": _format_service_time(form.get("service_time")),
        "location": (form.get("location") or "").strip(),
        "presiding": (form.get("presiding") or "").strip(),
        "conducting": (form.get("conducting") or "").strip(),
        "welcome_text": (form.get("welcome_text") or "").strip(),
        "opening_hymn_line": opening["line"],
        "opening_hymn": opening,
        "invocation": (form.get("invocation") or "").strip(),
        "speaker_1": (form.get("speaker_1") or "").strip(),
        "speaker_1_topic": (form.get("speaker_1_topic") or "").strip(),
        "speaker_2": (form.get("speaker_2") or "").strip(),
        "speaker_2_topic": (form.get("speaker_2_topic") or "").strip(),
        "musical_number": (form.get("musical_number") or "").strip(),
        "candidate_name": candidate,
        "baptism_by": (form.get("baptism_by") or "").strip(),
        "confirmation_by": (form.get("confirmation_by") or "").strip(),
        "confirmation_text": confirmation_text,
        "closing_hymn_line": closing["line"],
        "closing_hymn": closing,
        "benediction": (form.get("benediction") or "").strip(),
        "reception_notes": (form.get("reception_notes") or "").strip(),
    }


def build_baptism_text(data: dict) -> str:
    lines = ["Baptismal Service", ""]

    if data.get("service_date_display"):
        line = data["service_date_display"]
        if data.get("service_time_display"):
            line += f" at {data['service_time_display']}"
        lines.append(line)
    if data.get("location"):
        lines.append(data["location"])
    lines.append("")

    for label, key in (
        ("Presiding", "presiding"),
        ("Conducting", "conducting"),
    ):
        if data.get(key):
            lines.append(f"{label}: {data[key]}")

    if data.get("welcome_text"):
        lines.append("")
        lines.append(data["welcome_text"])

    lines.append("")
    if data.get("opening_hymn_line"):
        lines.append(f"Opening Hymn: {data['opening_hymn_line']}")
    if data.get("invocation"):
        lines.append(f"Invocation: {data['invocation']}")

    lines.append("")
    if data.get("speaker_1"):
        talk = data["speaker_1"]
        if data.get("speaker_1_topic"):
            talk += f" — {data['speaker_1_topic']}"
        lines.append(f"Talk: {talk}")
    if data.get("speaker_2"):
        talk = data["speaker_2"]
        if data.get("speaker_2_topic"):
            talk += f" — {data['speaker_2_topic']}"
        lines.append(f"Talk: {talk}")
    if data.get("musical_number"):
        lines.append(f"Musical number: {data['musical_number']}")

    lines.append("")
    if data.get("candidate_name"):
        lines.append(f"Baptism of {data['candidate_name']}")
    if data.get("baptism_by"):
        lines.append(f"Baptism performed by {data['baptism_by']}")

    lines.append("")
    confirmation = resolve_confirmation_text(data.get("candidate_name"), data.get("confirmation_text"))
    if confirmation:
        lines.append(confirmation)
    if data.get("confirmation_by"):
        lines.append(f"Confirmation by {data['confirmation_by']}")

    lines.append("")
    if data.get("closing_hymn_line"):
        lines.append(f"Closing Hymn: {data['closing_hymn_line']}")
    if data.get("benediction"):
        lines.append(f"Benediction: {data['benediction']}")

    if data.get("reception_notes"):
        lines.append("")
        lines.append(data["reception_notes"])

    return "\n".join(lines).strip() + "\n"


def export_docx(data: dict) -> bytes:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    def set_landscape(section) -> None:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    def emu_to_twips(emu: int) -> int:
        return int(emu / 635)

    def set_table_width(table, width_emu) -> None:
        table.autofit = False
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:w"), str(emu_to_twips(width_emu)))
        tbl_w.set(qn("w:type"), "dxa")
        tbl_pr.append(tbl_w)
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")

    def set_cell_width(cell, width_emu) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_w = OxmlElement("w:tcW")
        tc_w.set(qn("w:w"), str(emu_to_twips(width_emu)))
        tc_w.set(qn("w:type"), "dxa")
        tc_pr.append(tc_w)

    def set_cell_margins(cell, *, top=100, bottom=100, start=120, end=120) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_mar = tc_pr.find(qn("w:tcMar"))
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for edge, val in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
            node = tc_mar.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(val))
            node.set(qn("w:type"), "dxa")

    def set_row_page_break_before(row) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:pageBreakBefore")) is None:
            tr_pr.append(OxmlElement("w:pageBreakBefore"))

    doc = Document()
    section = doc.sections[0]
    set_landscape(section)
    margin = Inches(0.45)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    content_width = section.page_width - section.left_margin - section.right_margin
    panel_width = content_width // 2
    body_font = "Times New Roman"
    program_size = Pt(11)
    lyric_size = Pt(7.5)
    cover_title_size = Pt(18)
    cover_sub_size = Pt(13)

    opening = dict(data.get("opening_hymn") or {})
    closing = dict(data.get("closing_hymn") or {})
    for hymn in (opening, closing):
        if hymn.get("number") and not (hymn.get("lyrics") or "").strip():
            hymn["lyrics"] = hymn_lyrics(hymn.get("number"), hymn.get("book") or HYMN_BOOK_CHILDREN)

    def clear_cell(cell) -> None:
        cell.text = ""

    def add_run(paragraph, text: str, *, bold: bool = False, size: Pt | None = None, center: bool = False):
        if center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = body_font
        if size:
            run.font.size = size
        return run

    def add_para(
        cell,
        text: str = "",
        *,
        bold: bool = False,
        center: bool = False,
        size: Pt | None = program_size,
        after: float = 6,
        before: float = 0,
        leading: float = 1.18,
    ):
        p = cell.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = leading
        if text:
            add_run(p, text, bold=bold, center=center, size=size)
        elif center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def add_labeled_para(cell, label: str, value: str, *, after: float = 8, before: float = 0) -> None:
        if not value:
            return
        p = cell.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.18
        label_run = add_run(p, label + ": ", bold=True, size=program_size)
        label_run.font.name = body_font
        add_run(p, value, size=program_size)

    def add_multiline(
        cell,
        text: str,
        *,
        size: Pt | None = program_size,
        after_last: float = 8,
        line_after: float = 4,
        leading: float = 1.18,
        center: bool = False,
    ) -> None:
        parts = [part.strip() for part in (text or "").splitlines() if part.strip()]
        for i, part in enumerate(parts):
            add_para(
                cell,
                part,
                size=size,
                after=after_last if i == len(parts) - 1 else line_after,
                leading=leading,
                center=center,
            )

    def add_centered_block(cell, text: str, *, size: Pt | None = program_size, after: float = 6) -> None:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(after)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = body_font
        run.font.size = size

    def hymn_panel_heading(hymn: dict, label: str) -> str:
        parts = [label]
        if hymn.get("number"):
            parts.append(f"#{hymn['number']}")
        if hymn.get("title"):
            parts.append(hymn["title"])
        if hymn.get("book") == HYMN_BOOK_CHILDREN:
            parts.append("(Children's Songbook)")
        elif hymn.get("book") == HYMN_BOOK_HYMNS and hymn.get("title"):
            parts.append("(Hymns)")
        return "  ".join(parts)

    def fill_hymn_lyrics_panel(cell, hymn: dict, label: str) -> None:
        clear_cell(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_margins(cell, top=80, bottom=80, start=100, end=100)
        number = hymn.get("number")
        book = hymn.get("book") or HYMN_BOOK_CHILDREN
        lyrics = (hymn.get("lyrics") or "").strip() or hymn_lyrics(number, book)
        add_para(cell, hymn_panel_heading(hymn, label), bold=True, center=True, size=Pt(10), after=8, leading=1.1)
        if lyrics:
            add_multiline(cell, lyrics, size=lyric_size, after_last=0, line_after=1, leading=1.05)
        elif book == HYMN_BOOK_CHILDREN and number:
            add_para(
                cell,
                f"Children's Songbook #{number} — see songbook for lyrics.",
                center=True,
                size=lyric_size,
                after=0,
            )
        elif hymn.get("title"):
            add_para(cell, "Sing from the hymnbook.", center=True, size=lyric_size, after=0)

    def fill_program_panel(cell) -> None:
        clear_cell(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_margins(cell, top=110, bottom=90, start=130, end=110)

        add_para(cell, "Order of Service", bold=True, center=True, size=Pt(14), after=12, before=2, leading=1.15)

        if data.get("presiding"):
            add_labeled_para(cell, "Presiding", data["presiding"], after=9)
        if data.get("conducting"):
            add_labeled_para(cell, "Conducting", data["conducting"], after=9)
        if data.get("welcome_text"):
            add_multiline(cell, data["welcome_text"], after_last=10, line_after=5, leading=1.22)

        if data.get("opening_hymn_line"):
            add_labeled_para(cell, "Opening hymn", data["opening_hymn_line"], after=9, before=2)
        if data.get("invocation"):
            add_labeled_para(cell, "Invocation", data["invocation"], after=10)

        if data.get("speaker_1"):
            talk = data["speaker_1"]
            if data.get("speaker_1_topic"):
                talk += f" — {data['speaker_1_topic']}"
            add_labeled_para(cell, "Talk", talk, after=9, before=2)
        if data.get("speaker_2"):
            talk = data["speaker_2"]
            if data.get("speaker_2_topic"):
                talk += f" — {data['speaker_2_topic']}"
            add_labeled_para(cell, "Talk", talk, after=9)
        if data.get("musical_number"):
            add_labeled_para(cell, "Musical number", data["musical_number"], after=10)

        if data.get("candidate_name"):
            add_para(cell, f"Baptism of {data['candidate_name']}", bold=True, after=7, before=3)
        if data.get("baptism_by"):
            add_labeled_para(cell, "Baptism by", data["baptism_by"], after=9)

        confirmation = resolve_confirmation_text(data.get("candidate_name"), data.get("confirmation_text"))
        if confirmation:
            add_multiline(cell, confirmation, after_last=9, line_after=5, leading=1.22)

        if data.get("confirmation_by"):
            add_labeled_para(cell, "Confirmation by", data["confirmation_by"], after=9, before=2)

        if data.get("closing_hymn_line"):
            add_labeled_para(cell, "Closing hymn", data["closing_hymn_line"], after=9)
        if data.get("benediction"):
            add_labeled_para(cell, "Benediction", data["benediction"], after=7)
        if data.get("reception_notes"):
            add_multiline(cell, data["reception_notes"], after_last=0, line_after=5, leading=1.18)

    def fill_cover_panel(cell) -> None:
        clear_cell(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=70, bottom=70, start=90, end=90)
        if _COVER_IMAGE_PATH.is_file():
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(_COVER_IMAGE_PATH), width=Inches(3.1))
            p.paragraph_format.space_after = Pt(10)

        add_para(cell, "Baptismal Service", bold=True, center=True, size=cover_title_size, after=8, leading=1.1)
        if data.get("candidate_name"):
            add_para(cell, data["candidate_name"], bold=True, center=True, size=cover_sub_size, after=8, leading=1.1)

        subtitle_lines = []
        if data.get("service_date_display"):
            subtitle_lines.append(data["service_date_display"])
        if data.get("service_time_display"):
            subtitle_lines.append(data["service_time_display"])
        if data.get("location"):
            subtitle_lines.append(data["location"])
        if subtitle_lines:
            add_centered_block(cell, "\n".join(subtitle_lines), size=Pt(11), after=0)

    # One 2x2 table: row 1 = outside (lyrics | cover), row 2 = inside (program | lyrics).
    # Page break before row 2 avoids extra blank pages from standalone break paragraphs.
    table = doc.add_table(rows=2, cols=2)
    set_table_width(table, content_width)

    row_outside = table.rows[0]
    row_inside = table.rows[1]
    set_row_page_break_before(row_inside)

    for cell, width in (
        (row_outside.cells[0], panel_width),
        (row_outside.cells[1], panel_width),
        (row_inside.cells[0], panel_width),
        (row_inside.cells[1], panel_width),
    ):
        set_cell_width(cell, width)

    fill_hymn_lyrics_panel(row_outside.cells[0], closing, "Closing Hymn")
    fill_cover_panel(row_outside.cells[1])
    fill_program_panel(row_inside.cells[0])
    fill_hymn_lyrics_panel(row_inside.cells[1], opening, "Opening Hymn")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
