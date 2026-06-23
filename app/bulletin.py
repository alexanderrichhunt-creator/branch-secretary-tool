from __future__ import annotations

import io
import re
from datetime import date, datetime, timedelta

from .hymns import hymn_display, hymn_title

URL_PATTERN = re.compile(r"https?://[^\s<>\"']+")


def _split_trailing_url_punctuation(url: str) -> tuple[str, str]:
    trailing = ""
    while url and url[-1] in ".,);":
        trailing = url[-1] + trailing
        url = url[:-1]
    return url, trailing

DEFAULT_BULLETIN = {
    "presiding": "Michael Reynolds, Madisonville Branch President",
    "conducting": "Pat Gaume, 1st Counselor",
    "on_the_stand": "Will Ross, 2nd Counselor, Madisonville Branch Presidency",
    "welcome_text": (
        "Welcome… any visitors who might be in attendance. "
        "Acknowledgment of any Stake visitors."
    ),
    "opening_hymn_num": "6",
    "opening_hymn_title": "",
    "invocation": "(by invitation)",
    "branch_business": (
        "Seminary Graduation on Sunday, May 17th, 5:00 p.m. at the Creighton Building in Conroe."
    ),
    "stake_business": "(if any)",
    "announcements": (
        "Here is the Zoom Link for the Madisonville Branch. It should activate a few minutes "
        "prior to 10:00 a.m. on Sundays. Copy and paste the Link into your Zoom App.\n"
        "https://zoom.us/j/92713005551\n"
        "Check for other news items at the Madisonville Branch Page on Facebook."
    ),
    "sacrament_notes": (
        "After the singing of the Sacrament Hymn, the Zoom broadcast portion of this meeting "
        "will be turned off; however, it will be turned back on after the Sacrament has been "
        "passed. If you have been authorized to bless the Sacrament during your meeting at home, "
        "please take this as your opportunity to do so."
    ),
    "sacrament_hymn_num": "190",
    "sacrament_hymn_title": "",
    "intermediate_hymn_num": "",
    "intermediate_hymn_title": "",
    "closing_hymn_num": "141",
    "closing_hymn_title": "",
    "benediction": "(by invitation)",
}

# Fields stored as branch defaults (not meeting_date or speakers_text).
SAVABLE_BULLETIN_KEYS = (
    "presiding",
    "conducting",
    "on_the_stand",
    "welcome_text",
    "opening_hymn_num",
    "opening_hymn_title",
    "invocation",
    "stake_business",
    "announcements",
    "sacrament_notes",
    "sacrament_hymn_num",
    "sacrament_hymn_title",
    "intermediate_hymn_num",
    "intermediate_hymn_title",
    "closing_hymn_num",
    "closing_hymn_title",
    "benediction",
)


def get_branch_bulletin_defaults() -> dict:
    """Built-in template merged with saved branch defaults from the database."""
    from .models import BulletinDefaults

    merged = dict(DEFAULT_BULLETIN)
    row = BulletinDefaults.query.get(1)
    if not row:
        return merged
    for key in SAVABLE_BULLETIN_KEYS:
        merged[key] = getattr(row, key) or ""
    return merged


def has_saved_branch_defaults() -> bool:
    from .models import BulletinDefaults

    return BulletinDefaults.query.get(1) is not None


def save_branch_bulletin_defaults(form) -> None:
    from . import db
    from .models import BulletinDefaults

    row = BulletinDefaults.query.get(1)
    if not row:
        row = BulletinDefaults(id=1)
        db.session.add(row)
    for key in SAVABLE_BULLETIN_KEYS:
        setattr(row, key, (form.get(key) or "").strip())
    row.updated_at = datetime.utcnow()
    db.session.commit()


def default_sacrament_sunday(today: date | None = None) -> date:
    today = today or date.today()
    days_until = (6 - today.weekday()) % 7
    if days_until == 0:
        return today
    return today + timedelta(days=days_until)


def _parse_hymn_num(raw: str | None) -> int | None:
    raw = (raw or "").strip()
    if not raw:
        return None
    try:
        n = int(raw)
        return n if n > 0 else None
    except ValueError:
        return None


def _format_meeting_date(d: date | str | None) -> str:
    if isinstance(d, str):
        try:
            d = datetime.strptime(d, "%Y-%m-%d").date()
        except ValueError:
            return d
    if not d:
        return ""
    return d.strftime("%B %d, %Y").replace(" 0", " ")


def bulletin_from_form(form) -> dict:
    meeting_date_raw = (form.get("meeting_date") or "").strip()
    meeting_date = None
    if meeting_date_raw:
        try:
            meeting_date = datetime.strptime(meeting_date_raw, "%Y-%m-%d").date()
        except ValueError:
            meeting_date = None

    opening_num = (form.get("opening_hymn_num") or "").strip()
    opening_title = (form.get("opening_hymn_title") or "").strip()
    sacrament_num = (form.get("sacrament_hymn_num") or "").strip()
    sacrament_title = (form.get("sacrament_hymn_title") or "").strip()
    intermediate_num = (form.get("intermediate_hymn_num") or "").strip()
    intermediate_title = (form.get("intermediate_hymn_title") or "").strip()
    closing_num = (form.get("closing_hymn_num") or "").strip()
    closing_title = (form.get("closing_hymn_title") or "").strip()

    return {
        "meeting_date": meeting_date,
        "meeting_date_display": _format_meeting_date(meeting_date or meeting_date_raw),
        "presiding": (form.get("presiding") or "").strip(),
        "conducting": (form.get("conducting") or "").strip(),
        "on_the_stand": (form.get("on_the_stand") or "").strip(),
        "welcome_text": (form.get("welcome_text") or "").strip(),
        "opening_hymn_num": opening_num,
        "opening_hymn_title": opening_title,
        "opening_hymn_line": hymn_display(opening_num, opening_title),
        "invocation": (form.get("invocation") or "").strip(),
        "branch_business": (form.get("branch_business") or "").strip(),
        "stake_business": (form.get("stake_business") or "").strip(),
        "announcements": (form.get("announcements") or "").strip(),
        "sacrament_notes": (form.get("sacrament_notes") or "").strip(),
        "sacrament_hymn_num": sacrament_num,
        "sacrament_hymn_title": sacrament_title,
        "sacrament_hymn_line": hymn_display(sacrament_num, sacrament_title),
        "intermediate_hymn_num": intermediate_num,
        "intermediate_hymn_title": intermediate_title,
        "intermediate_hymn_line": hymn_display(intermediate_num, intermediate_title),
        "speakers_text": (form.get("speakers_text") or "").strip(),
        "speakers_mode": (form.get("speakers_mode") or SPEAKERS_MODE_TALKS).strip(),
        "closing_hymn_num": closing_num,
        "closing_hymn_title": closing_title,
        "closing_hymn_line": hymn_display(closing_num, closing_title),
        "benediction": (form.get("benediction") or "").strip(),
    }


def build_bulletin_text(data: dict, talks=None) -> str:
    lines = [
        "Sacrament Meeting",
        data.get("meeting_date_display") or "",
        "",
    ]
    if data.get("presiding"):
        lines.append(f"Presiding: {data['presiding']}")
    if data.get("conducting"):
        lines.append(f"Conducting: {data['conducting']}")
    if data.get("on_the_stand"):
        lines.append(f"On the stand: {data['on_the_stand']}")
    lines.append("")
    if data.get("welcome_text"):
        lines.append(data["welcome_text"])
        lines.append("")
    if data.get("opening_hymn_line"):
        lines.append(f"Opening Hymn: {data['opening_hymn_line']}")
    if data.get("invocation"):
        lines.append(f"Invocation: {data['invocation']}")
    lines.append("")
    lines.append("Branch Business:")
    lines.append(data.get("branch_business") or "")
    lines.append("")
    lines.append(f"Stake Business: {data.get('stake_business') or ''}")
    lines.append("")
    if data.get("announcements"):
        lines.append(data["announcements"])
        lines.append("")
    if data.get("sacrament_notes"):
        lines.append(data["sacrament_notes"])
        lines.append("")
    if data.get("sacrament_hymn_line"):
        lines.append(f"The Sacrament Hymn is {data['sacrament_hymn_line']}")
    lines.append("")
    lines.extend(program_lines_after_sacrament(data, talks))
    if data.get("closing_hymn_line"):
        lines.append(f"Closing Hymn {data['closing_hymn_line']}")
    if data.get("benediction"):
        lines.append(f"Benediction: {data['benediction']}")
    return "\n".join(lines).strip() + "\n"


def _iter_bulletin_lines(data: dict) -> list[str]:
    return build_bulletin_text(data).split("\n")


def export_docx(data: dict, talks=None) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Inches, Pt

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    body_size = 11
    body_after = 4
    section_after = 8
    leading = 1.22

    def set_para_spacing(pf, *, after: float = body_after, leading_val: float = leading) -> None:
        pf.space_before = Pt(0)
        pf.space_after = Pt(after)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = leading_val

    def add_line(
        text: str,
        *,
        bold: bool = False,
        size: float = body_size,
        center: bool = False,
        after: float = body_after,
        leading_val: float = leading,
    ) -> None:
        p = doc.add_paragraph()
        set_para_spacing(p.paragraph_format, after=after, leading_val=leading_val)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)

    def add_hyperlink(paragraph, text: str, url: str) -> None:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml.shared import OxmlElement, qn

        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        r_pr.append(color)

        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Times New Roman")
        r_fonts.set(qn("w:hAnsi"), "Times New Roman")
        r_pr.append(r_fonts)

        size_el = OxmlElement("w:sz")
        size_el.set(qn("w:val"), str(int(body_size * 2)))
        r_pr.append(size_el)

        new_run.append(r_pr)
        text_el = OxmlElement("w:t")
        text_el.text = text
        new_run.append(text_el)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def add_paragraph_with_links(
        text: str,
        *,
        after: float = body_after,
        bold: bool = False,
        center: bool = False,
        size: float = body_size,
        leading_val: float = leading,
    ) -> None:
        p = doc.add_paragraph()
        set_para_spacing(p.paragraph_format, after=after, leading_val=leading_val)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        pos = 0
        for match in URL_PATTERN.finditer(text):
            if match.start() > pos:
                run = p.add_run(text[pos : match.start()])
                run.bold = bold
                run.font.name = "Times New Roman"
                run.font.size = Pt(size)
            url = match.group(0)
            url, trailing = _split_trailing_url_punctuation(url)
            add_hyperlink(p, url, url)
            if trailing:
                run = p.add_run(trailing)
                run.bold = bold
                run.font.name = "Times New Roman"
                run.font.size = Pt(size)
            pos = match.end()

        if pos < len(text):
            run = p.add_run(text[pos:])
            run.bold = bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)

    def add_multiline(text: str, *, after_last: float = body_after, linkify: bool = True, **kwargs) -> None:
        parts = [part.strip() for part in (text or "").splitlines() if part.strip()]
        for i, part in enumerate(parts):
            is_last = i == len(parts) - 1
            after = after_last if is_last else body_after
            if linkify and URL_PATTERN.search(part):
                add_paragraph_with_links(part, after=after, **kwargs)
            else:
                add_line(part, after=after, **kwargs)

    def add_labeled_line(
        label: str,
        value: str,
        *,
        separator: str = ": ",
        after: float = body_after,
        show_empty: bool = False,
    ) -> None:
        if not value and not show_empty:
            return
        p = doc.add_paragraph()
        set_para_spacing(p.paragraph_format, after=after, leading_val=leading)
        label_run = p.add_run(label + separator)
        label_run.bold = True
        label_run.font.name = "Times New Roman"
        label_run.font.size = Pt(body_size)
        if value:
            value_run = p.add_run(value)
            value_run.font.name = "Times New Roman"
            value_run.font.size = Pt(body_size)

    add_line("Sacrament Meeting", bold=True, size=15, center=True, after=3, leading_val=1.0)
    if data.get("meeting_date_display"):
        add_line(
            data["meeting_date_display"],
            bold=True,
            size=11,
            center=True,
            after=12,
            leading_val=1.0,
        )

    add_labeled_line("Presiding", data.get("presiding") or "", after=body_after)
    add_labeled_line("Conducting", data.get("conducting") or "", after=body_after)
    add_labeled_line("On the stand", data.get("on_the_stand") or "", after=section_after)

    if data.get("welcome_text"):
        add_multiline(data["welcome_text"], after_last=section_after)

    add_labeled_line("Opening Hymn", data.get("opening_hymn_line") or "", after=body_after)
    add_labeled_line("Invocation", data.get("invocation") or "", after=section_after)

    add_line("Branch Business:", bold=True, after=2)
    add_multiline(data.get("branch_business") or "", after_last=section_after)

    add_labeled_line(
        "Stake Business",
        data.get("stake_business") or "",
        after=section_after,
        show_empty=True,
    )

    if data.get("announcements"):
        add_multiline(data["announcements"], after_last=section_after)

    if data.get("sacrament_notes"):
        add_multiline(data["sacrament_notes"], after_last=section_after)

    if data.get("sacrament_hymn_line"):
        add_labeled_line(
            "The Sacrament Hymn is",
            data["sacrament_hymn_line"],
            separator=" ",
            after=section_after,
        )

    for program_line in program_lines_after_sacrament(data, talks):
        if program_line.startswith("Intermediate Hymn:"):
            add_labeled_line(
                "Intermediate Hymn",
                program_line.split(":", 1)[1].strip(),
                after=section_after,
            )
        elif program_line:
            add_multiline(program_line, after_last=section_after)

    add_labeled_line("Closing Hymn", data.get("closing_hymn_line") or "", separator=" ", after=body_after)
    add_labeled_line("Benediction", data.get("benediction") or "", after=body_after)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def speakers_text_for_talks(talks) -> str:
    talks = sort_assigned_talks(talks)
    names = [_talk_display_name(t) for t in talks]
    names = [n for n in names if n and n != "—"]
    if not names:
        return ""
    if len(names) == 1:
        return f"Our speaker today will be {names[0]}."
    return "Our speakers today will be " + " followed by ".join(names) + "."


def speakers_text_blocks_for_talks(talks) -> list[str]:
    """Split speaker intro for bulletin layout with an intermediate hymn between groups."""
    talks = sort_assigned_talks(talks)
    names = [_talk_display_name(t) for t in talks]
    names = [n for n in names if n and n != "—"]
    if not names:
        return []
    if len(names) == 1:
        return [f"Our speaker today will be {names[0]}."]
    first = f"Our speaker today will be {names[0]}."
    rest = "Followed by " + " followed by ".join(names[1:]) + "."
    return [first, rest]


def speakers_text_for_talks_layout(talks, *, split_for_intermediate: bool = False) -> str:
    talks = sort_assigned_talks(talks)
    if split_for_intermediate and len(talks) >= 2:
        return "\n\n".join(speakers_text_blocks_for_talks(talks))
    return speakers_text_for_talks(talks)


def talk_sort_key(talk) -> tuple:
    order = getattr(talk, "sort_order", 0) or 0
    if order <= 0:
        order = 999
    talk_date = getattr(talk, "talk_date", None)
    talk_id = getattr(talk, "id", 0) or 0
    return (order, talk_date or date.min, talk_id)


def sort_assigned_talks(talks) -> list:
    return sorted(regular_assigned_talks(talks), key=talk_sort_key)


def regular_assigned_talks(talks) -> list:
    return [t for t in talks if not is_special_meeting_talk(t)]


def has_intermediate_hymn(data: dict) -> bool:
    return bool((data.get("intermediate_hymn_line") or "").strip())


def program_lines_after_sacrament(data: dict, talks=None) -> list[str]:
    """Speaker paragraphs with optional intermediate hymn between first and remaining speakers."""
    lines: list[str] = []
    intermediate = (data.get("intermediate_hymn_line") or "").strip()
    speakers_mode = (data.get("speakers_mode") or SPEAKERS_MODE_TALKS).strip()
    speakers_text = (data.get("speakers_text") or "").strip()
    assigned = sort_assigned_talks(talks or [])

    if _special_meeting_mode_from_data(speakers_mode, talks):
        text = speakers_text or _special_meeting_speakers_text(speakers_mode, talks)
        if text:
            lines.append(text)
            lines.append("")
        return lines

    if assigned and intermediate and len(assigned) >= 2:
        blocks = speakers_text_blocks_for_talks(assigned)
        if blocks:
            lines.append(blocks[0])
            lines.append("")
            lines.append(f"Intermediate Hymn: {intermediate}")
            lines.append("")
            if len(blocks) > 1:
                lines.append(blocks[1])
                lines.append("")
        return lines

    if assigned:
        text = speakers_text or speakers_text_for_talks(assigned)
        if text:
            lines.append(text)
            lines.append("")
        if intermediate:
            lines.append(f"Intermediate Hymn: {intermediate}")
            lines.append("")
        return lines

    if speakers_text:
        parts = [part.strip() for part in speakers_text.split("\n\n") if part.strip()]
        if intermediate and len(parts) >= 2:
            lines.append(parts[0])
            lines.append("")
            lines.append(f"Intermediate Hymn: {intermediate}")
            lines.append("")
            lines.extend(parts[1:])
            lines.append("")
        else:
            lines.append(speakers_text)
            lines.append("")
            if intermediate:
                lines.append(f"Intermediate Hymn: {intermediate}")
                lines.append("")
    elif intermediate:
        lines.append(f"Intermediate Hymn: {intermediate}")
        lines.append("")

    return lines


SPEAKERS_MODE_TALKS = "talks"
SPEAKERS_MODE_FAST_TESTIMONY = "fast_testimony"
SPEAKERS_MODE_BRANCH_CONFERENCE = "branch_conference"
SPEAKERS_MODE_STAKE_CONFERENCE = "stake_conference"
SPEAKERS_MODE_GENERAL_CONFERENCE = "general_conference"

TALK_KIND_ASSIGNED = "assigned"
TALK_KIND_FAST_TESTIMONY = "fast_testimony"
TALK_KIND_BRANCH_CONFERENCE = "branch_conference"
TALK_KIND_STAKE_CONFERENCE = "stake_conference"
TALK_KIND_GENERAL_CONFERENCE = "general_conference"

SPECIAL_MEETINGS = {
    TALK_KIND_FAST_TESTIMONY: {
        "label": "Fast and Testimony Meeting",
        "short_label": "Fast & Testimony Meeting",
        "speakers_text": "Today we will hold a Fast and Testimony Meeting.",
        "speakers_mode": SPEAKERS_MODE_FAST_TESTIMONY,
        "calendar_kind": "fast_testimony",
        "calendar_label": "Fast & Testimony Meeting",
    },
    TALK_KIND_BRANCH_CONFERENCE: {
        "label": "Branch Conference",
        "short_label": "Branch Conference",
        "speakers_text": "Today we will hold Branch Conference.",
        "speakers_mode": SPEAKERS_MODE_BRANCH_CONFERENCE,
        "calendar_kind": "branch_conference",
        "calendar_label": "Branch Conference",
    },
    TALK_KIND_STAKE_CONFERENCE: {
        "label": "Stake Conference",
        "short_label": "Stake Conference",
        "speakers_text": "Today we will hold Stake Conference.",
        "speakers_mode": SPEAKERS_MODE_STAKE_CONFERENCE,
        "calendar_kind": "stake_conference",
        "calendar_label": "Stake Conference",
    },
    TALK_KIND_GENERAL_CONFERENCE: {
        "label": "General Conference",
        "short_label": "General Conference",
        "speakers_text": "Today we will view General Conference.",
        "speakers_mode": SPEAKERS_MODE_GENERAL_CONFERENCE,
        "calendar_kind": "general_conference",
        "calendar_label": "General Conference",
    },
}

SPECIAL_TALK_KINDS = frozenset(SPECIAL_MEETINGS.keys())
SPECIAL_SPEAKERS_MODES = frozenset(meta["speakers_mode"] for meta in SPECIAL_MEETINGS.values())

FAST_TESTIMONY_LABEL = SPECIAL_MEETINGS[TALK_KIND_FAST_TESTIMONY]["label"]
FAST_TESTIMONY_SPEAKERS_TEXT = SPECIAL_MEETINGS[TALK_KIND_FAST_TESTIMONY]["speakers_text"]
BRANCH_CONFERENCE_LABEL = SPECIAL_MEETINGS[TALK_KIND_BRANCH_CONFERENCE]["label"]
STAKE_CONFERENCE_LABEL = SPECIAL_MEETINGS[TALK_KIND_STAKE_CONFERENCE]["label"]
GENERAL_CONFERENCE_LABEL = SPECIAL_MEETINGS[TALK_KIND_GENERAL_CONFERENCE]["label"]


def is_special_talk_kind(kind: str) -> bool:
    return kind in SPECIAL_TALK_KINDS


def label_for_talk_kind(kind: str) -> str | None:
    meta = SPECIAL_MEETINGS.get(kind)
    return meta["label"] if meta else None


def special_meeting_kind(talk) -> str | None:
    if getattr(talk, "member_id", None):
        return None
    speaker = (getattr(talk, "speaker_text", None) or "").strip().casefold()
    for kind, meta in SPECIAL_MEETINGS.items():
        if speaker == meta["label"].casefold():
            return kind
    return None


def is_special_meeting_talk(talk) -> bool:
    return special_meeting_kind(talk) is not None


def is_fast_testimony_talk(talk) -> bool:
    return special_meeting_kind(talk) == TALK_KIND_FAST_TESTIMONY


def is_branch_conference_talk(talk) -> bool:
    return special_meeting_kind(talk) == TALK_KIND_BRANCH_CONFERENCE


def is_stake_conference_talk(talk) -> bool:
    return special_meeting_kind(talk) == TALK_KIND_STAKE_CONFERENCE


def is_general_conference_talk(talk) -> bool:
    return special_meeting_kind(talk) == TALK_KIND_GENERAL_CONFERENCE


def _special_meeting_mode_from_data(speakers_mode: str, talks) -> bool:
    if speakers_mode in SPECIAL_SPEAKERS_MODES:
        return True
    return any(special_meeting_kind(t) for t in (talks or []))


def _special_meeting_speakers_text(speakers_mode: str, talks) -> str:
    for kind, meta in SPECIAL_MEETINGS.items():
        if speakers_mode == meta["speakers_mode"] or any(
            special_meeting_kind(t) == kind for t in (talks or [])
        ):
            return meta["speakers_text"]
    return ""


def special_meeting_meta(kind: str | None) -> dict | None:
    if not kind:
        return None
    return SPECIAL_MEETINGS.get(kind)


def is_first_sacrament_sunday(d: date) -> bool:
    return d.weekday() == 6 and d.day <= 7


def default_speakers_mode(meeting_date: date, talks=None) -> str:
    if talks:
        for talk in talks:
            kind = special_meeting_kind(talk)
            if kind:
                return SPECIAL_MEETINGS[kind]["speakers_mode"]
    if is_first_sacrament_sunday(meeting_date):
        return SPEAKERS_MODE_FAST_TESTIMONY
    return SPEAKERS_MODE_TALKS


def speakers_text_for_mode(mode: str, talks, *, split_for_intermediate: bool = False) -> str:
    if mode in SPECIAL_SPEAKERS_MODES or any(is_special_meeting_talk(t) for t in talks):
        return _special_meeting_speakers_text(mode, talks)
    return speakers_text_for_talks_layout(talks, split_for_intermediate=split_for_intermediate)


def resolved_hymn_title(defaults: dict, num_key: str, title_key: str) -> str:
    saved = (defaults.get(title_key) or "").strip()
    if saved:
        return saved
    return hymn_title(_parse_hymn_num(defaults.get(num_key)))


def bulletin_person_name(name: str) -> str:
    """Render stored 'Last, First' names as 'First Last' on the bulletin."""
    name = " ".join((name or "").strip().split())
    if not name or "," not in name:
        return name
    last, _, first = name.partition(",")
    last = last.strip()
    first = first.strip()
    if first and last:
        return f"{first} {last}"
    return name


def _talk_display_name(t) -> str:
    if t.member_id and t.member is not None:
        return bulletin_person_name(t.member.full_name)
    return (getattr(t, "speaker_text", None) or "").strip() or "—"
